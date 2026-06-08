"""
인공지능실습 — 스마트팜 식물 병해 진단 AI 최종보고서 (.docx)
- AI 중심 구성
- 팀원: 봉준표, 이동제, 유광현, 최성옥
- 기존 보고서(스마트팜_연동앱_프로젝트_최종보고서.docx) 토대
- 앱 캡처(/Users/bongjunpyo/Desktop/app_picture/) 9장 + ERD 1장 삽입
- 한글 폰트: 맑은 고딕 → 안 되면 Apple SD Gothic Neo 적용 (양쪽 다 세팅)
"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

OUT = "/Users/bongjunpyo/smartfarm-upgrade/스마트팜_AI_병해진단_최종보고서.docx"
PIC = "/Users/bongjunpyo/Desktop/app_picture"
LEG = "/Users/bongjunpyo/smartfarm-upgrade/report_assets/legacy_docx_img"

# ====== 새 Document ======
d = Document()

# 페이지 여백
for section in d.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.2)
    section.right_margin  = Cm(2.2)

# ====== 폰트 기본값 ======
# macOS / Windows / LibreOffice 모두에서 깨지지 않도록 가장 호환성 높은 폰트로 통일
KFONT = "Apple SD Gothic Neo"
EFONT = "Apple SD Gothic Neo"   # 라틴 문자도 동일 폰트로 통일하여 fallback 깨짐 방지

def set_run_font(run, name=KFONT, size=11, bold=False,
                 color=None, italic=False):
    run.font.name = EFONT
    run.font.size = Pt(size)
    run.font.bold = bold
    # italic은 한글 폰트에서 fake italic 렌더링 깨짐 위험 → 사용 금지
    run.font.italic = False
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)

# 스타일 기본 폰트
styles = d.styles
normal = styles["Normal"]
normal.font.name = KFONT
normal.font.size = Pt(11)
rpr = normal.element.get_or_add_rPr()
rFonts = rpr.find(qn("w:rFonts"))
if rFonts is None:
    rFonts = OxmlElement("w:rFonts")
    rpr.append(rFonts)
rFonts.set(qn("w:eastAsia"), KFONT)
rFonts.set(qn("w:ascii"), KFONT)
rFonts.set(qn("w:hAnsi"), KFONT)
rFonts.set(qn("w:cs"), KFONT)

# ====== 헬퍼 함수 ======
def add_para(text="", size=11, bold=False, color=None, align=None,
             space_after=4, space_before=0, indent=None, line=1.5):
    p = d.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if indent is not None:
        p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.line_spacing = line
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, color=color)
    return p

def add_heading1(text):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, size=18, bold=True, color=(0x15, 0x80, 0x3d))
    # 밑줄 강조 (탑 보더)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:color"), "16A34A")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_heading2(text):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, size=14, bold=True, color=(0x0f, 0x17, 0x2a))
    return p

def add_heading3(text):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, size=12, bold=True, color=(0x15, 0x80, 0x3d))
    return p

def add_bullet(text, size=11):
    p = d.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.4
    r = p.add_run(text)
    set_run_font(r, size=size)

_NUM_COUNTER = {"n": 0}
def reset_numbering():
    _NUM_COUNTER["n"] = 0

def add_numbered(text, size=11):
    _NUM_COUNTER["n"] += 1
    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.8)
    r = p.add_run(f"{_NUM_COUNTER['n']}.  ")
    set_run_font(r, size=size, bold=True, color=(0x16, 0xa3, 0x4a))
    r = p.add_run(text)
    set_run_font(r, size=size)

def shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)

def add_table(rows, header=True, col_widths=None, header_fill="16A34A",
              alt_fill="F1F5F9"):
    tbl = d.add_table(rows=len(rows), cols=len(rows[0]))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            for r in tbl.rows:
                r.cells[i].width = Cm(w)
    for ri, row in enumerate(rows):
        for ci, txt in enumerate(row):
            cell = tbl.rows[ri].cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(str(txt))
            if ri == 0 and header:
                set_run_font(r, size=10, bold=True, color=(255, 255, 255))
                shade_cell(cell, header_fill)
            else:
                set_run_font(r, size=10)
                if alt_fill and ri % 2 == 0:
                    shade_cell(cell, alt_fill)
    # 표 뒤 간격
    add_para("", space_after=2, line=1.0)
    return tbl

def add_image(path, width_cm=12, caption=None):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Cm(width_cm))
    if caption:
        pc = d.add_paragraph()
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rc = pc.add_run(caption)
        set_run_font(rc, size=10, bold=True, color=(0x47, 0x55, 0x69))
        pc.paragraph_format.space_after = Pt(10)

def add_info_box(text, color=(0x16, 0xa3, 0x4a)):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:color"), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
    left.set(qn("w:space"), "10")
    pBdr.append(left)
    pPr.append(pBdr)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F0FDF4")
    pPr.append(shd)
    r = p.add_run(text)
    set_run_font(r, size=11, bold=False, color=(0x15, 0x80, 0x3d))

# ============================================================
# 표지
# ============================================================
p = d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
p.paragraph_format.space_after = Pt(8)
r = p.add_run("2026학년도 1학기")
set_run_font(r, size=12, color=(0x64, 0x74, 0x8b))

p = d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(40)
r = p.add_run("인공지능실습")
set_run_font(r, size=16, bold=True, color=(0x15, 0x80, 0x3d))

p = d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run("스마트팜 식물 병해 진단 AI")
set_run_font(r, size=28, bold=True, color=(0x0f, 0x17, 0x2a))

p = d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
r = p.add_run("최 종 보 고 서")
set_run_font(r, size=28, bold=True, color=(0x0f, 0x17, 0x2a))

p = d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
r = p.add_run("EfficientNet-B3 기반 식물 잎 23종 병해 분류 모델 학습")
set_run_font(r, size=13, bold=False, color=(0x16, 0xa3, 0x4a))

p = d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(50)
r = p.add_run("및 모바일 응용 시스템 구현")
set_run_font(r, size=13, bold=False, color=(0x16, 0xa3, 0x4a))

# 표지 정보 표
rows = [
    ["과 목 명",       "인공지능실습"],
    ["프로젝트 주제",   "EfficientNet-B3 기반 식물 병해 분류 AI 모델 학습 및 응용"],
    ["팀     원",      "봉준표 · 이동제 · 유광현 · 최성옥"],
    ["보고 기준일",     "2026년 6월 8일"],
    ["핵심 산출물",     "병해 분류 AI 모델(EfficientNet-B3 ONNX), 추론 API, 안드로이드 응용 앱"],
    ["검증 정확도",     "Val Accuracy 95.30% (23 클래스)"],
]
tbl = d.add_table(rows=len(rows), cols=2)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (k, v) in enumerate(rows):
    tbl.rows[i].cells[0].width = Cm(4.5)
    tbl.rows[i].cells[1].width = Cm(11.5)
    for ci, txt in enumerate([k, v]):
        c = tbl.rows[i].cells[ci]
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c.text = ""
        p = c.paragraphs[0]
        r = p.add_run(txt)
        set_run_font(r, size=11, bold=(ci == 0),
                     color=((255,255,255) if ci == 0 else (0x0f,0x17,0x2a)))
        if ci == 0:
            shade_cell(c, "16A34A")

d.add_page_break()

# ============================================================
# 목차
# ============================================================
add_heading1("목 차")
toc = [
    ("1.",   "프로젝트 개요"),
    ("2.",   "개발 목적 및 필요성"),
    ("3.",   "전체 시스템 구성"),
    ("4.",   "AI 모델 설계 및 학습 (★ 핵심)"),
    ("",     "    4.1  데이터셋 및 분류 클래스"),
    ("",     "    4.2  데이터 전처리 — HSV Anomaly Score"),
    ("",     "    4.3  모델 구조 — EfficientNet-B3"),
    ("",     "    4.4  학습 설정 및 학습 곡선"),
    ("",     "    4.5  학습 결과 및 평가"),
    ("",     "    4.6  ONNX 변환 및 추론 검증"),
    ("",     "    4.7  그리드 추론 알고리즘"),
    ("5.",   "AI 모델의 응용 — 모바일 앱 연동"),
    ("",     "    5.1  백엔드 추론 API (FastAPI)"),
    ("",     "    5.2  안드로이드 앱 화면 구성"),
    ("6.",   "데이터베이스 설계"),
    ("7.",   "구현 결과 및 검증"),
    ("8.",   "역할 분담"),
    ("9.",   "개발 과정에서 느낀 점"),
    ("10.",  "향후 개선 계획"),
    ("11.",  "결론"),
]
for n, t in toc:
    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.3
    if n:
        r = p.add_run(f"{n} ")
        set_run_font(r, size=12, bold=True, color=(0x16, 0xa3, 0x4a))
    r = p.add_run(t)
    set_run_font(r, size=12, color=(0x0f, 0x17, 0x2a))

d.add_page_break()

# ============================================================
# 1. 프로젝트 개요
# ============================================================
add_heading1("1. 프로젝트 개요")
add_para(
    "본 프로젝트는 인공지능실습 과목의 일환으로, 식물 잎 이미지에서 병해 유형을 자동으로 "
    "판별하는 합성곱 신경망(CNN) 모델을 직접 설계·학습·검증하고, 그 결과를 모바일 앱과 "
    "백엔드 API를 통해 실제 사용자에게 전달할 수 있는 형태로 응용하는 것을 목표로 한다. "
    "특히 단순한 데모 수준의 학습 실습을 넘어, 학습된 모델을 ONNX로 변환하여 서버 추론과 "
    "온디바이스 추론 모두를 지원하는 실전형 파이프라인을 구축하였다."
)
add_para(
    "본 보고서는 인공지능실습 과목의 평가 대상이므로 AI 모델의 설계 의사결정, 전처리 "
    "전략, 학습 결과, 평가 지표를 중심으로 서술하며, 모바일 앱과 백엔드는 학습된 AI "
    "모델을 어떻게 응용하였는지를 보여주는 보조 산출물로 다룬다."
)

add_info_box(
    "핵심 한 줄 요약 — Kaggle New Plant Diseases Dataset 87,000장으로 EfficientNet-B3를 "
    "학습하여 23개 식물 병해 클래스를 검증 정확도 95.30%로 분류하는 모델을 완성하였고, "
    "이를 ONNX로 변환해 FastAPI 백엔드 + 안드로이드 앱과 연동하여 실제 농가가 사용 가능한 "
    "수준의 진단 서비스를 구현하였다."
)

# 표지 옆 팀 정보 표
add_heading2("팀 구성")
team_rows = [
    ["이름",     "직책",     "주요 역할"],
    ["봉준표",   "팀장",     "AI 모델 학습(EfficientNet-B3) · Android–AI 연동 · 프로젝트 설계"],
    ["이동제",   "팀원",     "백엔드 개발 · 프론트엔드 수정 · DB 수정"],
    ["유광현",   "팀원",     "DB 구성 · FastAPI 연동"],
    ["최성옥",   "팀원",     "앱 개발 · 모니터링 화면 구현"],
]
add_table(team_rows, col_widths=[2.5, 2.0, 11.5])

# ============================================================
# 2. 개발 목적 및 필요성
# ============================================================
add_heading1("2. 개발 목적 및 필요성")
add_para(
    "비닐하우스·시설 재배 환경에서 식물 병해는 초기 발견이 늦어질수록 작물 전체로 빠르게 "
    "확산되며, 일반 농가가 잎의 미세한 색·반점 변화로 병해 종류를 식별하는 것은 매우 어렵다. "
    "본 프로젝트는 이 문제를 인공지능, 그중에서도 이미지 분류 모델로 해결할 수 있는지를 "
    "실증하고자 한다."
)
add_bullet("학습된 AI 모델이 23개 식물 병해 클래스를 사람보다 빠르고 일관되게 판별할 수 있는지를 검증한다.")
add_bullet("Kaggle 공개 데이터셋 기반의 학습 결과가 실제 농가 환경(촬영 환경 차이)에서도 응용 가능한지 확인한다.")
add_bullet("학습한 모델을 ONNX로 변환하여 서버·모바일 양쪽에서 추론 가능한 경량 모델로 배포하는 전 과정을 경험한다.")
add_bullet("AI 모델 출력(클래스, 확률)을 그대로 두지 않고, 모바일 앱의 위험도·알림·구역 상태 등으로 변환하는 응용 설계를 학습한다.")
add_bullet("데이터 전처리 — 모델 학습 — 평가 — 변환 — 추론 — UI 응용으로 이어지는 AI 파이프라인을 풀스택으로 구현한다.")

# ============================================================
# 3. 전체 시스템 구성
# ============================================================
add_heading1("3. 전체 시스템 구성")
add_para(
    "전체 시스템은 AI 모델을 중심에 두고, 모델의 입력(이미지)과 출력(진단 결과)을 양 끝의 "
    "안드로이드 앱과 PostgreSQL DB로 연결하는 구조이다. 모든 흐름이 학습된 EfficientNet-B3 "
    "모델의 추론 결과에 의해 결정된다는 점이 본 시스템의 핵심이다."
)
reset_numbering()
add_numbered("스마트팜 구역별 카메라가 주기적으로 식물 잎 이미지를 촬영한다.")
add_numbered("이미지는 FastAPI 백엔드의 /analyze 엔드포인트로 업로드되어 images 테이블에 저장된다.")
add_numbered("백엔드는 학습된 ONNX 모델을 onnxruntime으로 로드하여 클래스와 신뢰도를 산출한다.")
add_numbered("추론 결과는 analysis_results 테이블에 (구역, 이미지, 질병 유형, 신뢰도)로 기록된다.")
add_numbered("위험 임계치 초과 시 notifications 테이블에 알림이 생성되고 앱이 이를 표시한다.")
add_numbered("사용자는 구역 현황·사진 기록·알림·상세 진단 화면을 통해 결과를 확인한다.")

add_heading2("3.1 구성 요소")
rows = [
    ["구성 요소",          "주요 역할"],
    ["AI 모델",            "EfficientNet-B3 — 224×224 입력에서 23개 식물 병해 클래스를 분류"],
    ["전처리 모듈",        "HSV Anomaly Score 기반으로 병해 후보 70×70 타일을 추출하여 학습 입력 생성"],
    ["백엔드 API (FastAPI)", "이미지 업로드 수신, ONNX 추론 실행, 결과 저장, 알림 생성"],
    ["PostgreSQL DB",      "fields · images · analysis_results · notifications 4개 테이블로 결과 추적"],
    ["안드로이드 앱",       "구역 현황, 사진 기록, 알림, 구역 상세 화면으로 AI 결과를 시각화"],
]
add_table(rows, col_widths=[4.5, 13.5])

# ============================================================
# 4. AI 모델 설계 및 학습 (★)
# ============================================================
add_heading1("4. AI 모델 설계 및 학습 (★ 핵심)")
add_para(
    "본 장은 인공지능실습 과목의 핵심에 해당하며, 데이터 수집·전처리·모델 구조·학습 설정·"
    "평가·변환·추론 검증의 전 과정을 정리한다."
)

add_heading2("4.1 데이터셋 및 분류 클래스")
add_para(
    "Kaggle 공개 데이터셋인 New Plant Diseases Dataset(Vipoooool)을 사용하였다. 이 데이터셋은 "
    "총 약 87,000장의 식물 잎 이미지를 포함하며, 5종 작물(감자, 토마토, 피망, 사과, 옥수수)의 "
    "건강 상태와 병해 상태를 모두 포함하는 23개 클래스로 구성된다. 데이터는 kagglehub로 "
    "자동 다운로드하여 로컬 캐시(~/.cache/kagglehub/)에 저장한 뒤 학습에 사용하였다."
)
rows = [
    ["작물",        "클래스 수",  "포함 클래스"],
    ["감자",        "3",          "Early Blight, Late Blight, Healthy"],
    ["토마토",      "10",         "Bacterial Spot, Early Blight, Late Blight, Leaf Mold, Septoria, Spider Mites, Target Spot, YLCV, Mosaic, Healthy"],
    ["피망",        "2",          "Bacterial Spot, Healthy"],
    ["사과",        "4",          "Apple Scab, Black Rot, Cedar Apple Rust, Healthy"],
    ["옥수수",      "4",          "Cercospora, Common Rust, Northern Leaf Blight, Healthy"],
    ["합계",        "23",         "건강 5종 + 병해 18종"],
]
add_table(rows, col_widths=[3.0, 2.5, 12.5])

add_heading2("4.2 데이터 전처리 — HSV Anomaly Score")
add_para(
    "원본 이미지를 그대로 모델에 입력하면 잎이 차지하는 영역이 일부에 그치고 배경이 정확도를 "
    "떨어뜨리는 문제가 있었다. 이를 해결하기 위해 두 가지 전처리 전략을 비교 실험하였다."
)
add_heading3("1차 시도 — YOLOv8 기반 병반 탐지")
add_para(
    "YOLOv8s를 학습하여 병반 영역을 객체로 탐지한 뒤 해당 영역만 잘라 분류기에 입력하는 방식을 "
    "시도하였다. 학습 자체는 mAP50=0.994로 매우 좋게 나왔으나, ① 모바일 환경에서 두 모델(YOLO + "
    "EfficientNet)을 동시에 로드해야 하는 부담, ② YOLO 추론 시간으로 인한 응답 지연, ③ 병반 "
    "박스 라벨이 별도로 필요하여 학습 데이터를 추가 수집해야 하는 비용 때문에 폐기하였다."
)
add_heading3("2차 채택 — HSV Anomaly Score (최종)")
add_para(
    "건강한 잎의 HSV 색공간 평균값(Hue ≈ 75, 녹색)을 기준으로 각 타일이 얼마나 \"비정상\"인지를 "
    "수치화하는 방식이다. 세 가지 지표를 가중 합산하여 Anomaly Score를 산출한다."
)
add_info_box(
    "Anomaly Score = mean(|H − 75|) × 1.0  +  √Var(S) × 0.4  +  √Var(Laplacian)·clip(120) × 0.15"
)
add_para("처리 흐름은 다음과 같다.")
reset_numbering()
add_numbered("원본 이미지를 128×128 슬라이딩 윈도우(overlap 25%)로 분할한다.")
add_numbered("HSV 기반 배경 필터로 채도/명도 극단 픽셀이 30%를 넘는 타일을 제거한다.")
add_numbered("남은 타일별로 위 수식의 Anomaly Score를 계산한다.")
add_numbered("최고 점수 타일을 70×70으로 크롭하여 학습 캐시에 저장한다.")
add_numbered("학습 입력 단계에서 224×224로 BICUBIC 업스케일하여 EfficientNet-B3에 투입한다.")

add_para(
    "이 방식은 YOLO 의존성을 완전히 제거하면서도 학습 데이터에서 \"잎 + 병반 의심 영역\"만을 "
    "선별적으로 사용할 수 있게 해주었고, 결과적으로 모델 크기를 절반으로 줄이면서 정확도는 "
    "약 5%p 향상시키는 효과를 얻었다."
)

add_heading2("4.3 모델 구조 — EfficientNet-B3")
add_para(
    "EfficientNet-B3은 Compound Scaling으로 깊이·너비·해상도를 동시에 확장한 경량 CNN으로, "
    "ImageNet 사전학습 가중치를 활용할 수 있어 23개 클래스 분류에 적합하다고 판단하였다. "
    "본 프로젝트에서는 다음과 같이 분류기(Classifier)만 교체하여 전이학습을 수행하였다."
)
add_info_box(
    "EfficientNet-B3 (ImageNet pretrained)\n"
    "    └─ classifier\n"
    "         ├─ Dropout(p = 0.3)\n"
    "         └─ Linear(in_features = 1536  →  out_features = 23)"
)

add_heading2("4.4 학습 설정 및 학습 곡선")
rows = [
    ["항목",              "설정값"],
    ["Optimizer",         "AdamW (lr = 1e-4, weight_decay = 0.01)"],
    ["Scheduler",         "CosineAnnealingLR (T_max = 10, eta_min = 1e-6)"],
    ["Loss Function",     "CrossEntropyLoss (클래스 불균형 역수 가중치 적용)"],
    ["Epochs",            "10"],
    ["Batch Size",        "32"],
    ["Train 샘플",        "20,000장 (HSV 전처리 통과 후 선별)"],
    ["Val 샘플",          "4,000장"],
    ["Input Size",        "224 × 224 (70×70 크롭 → BICUBIC 업스케일)"],
    ["Augmentation",      "RandomHorizontalFlip(0.5), VerticalFlip(0.2), Rotation(15°), ColorJitter, ImageNet Normalize"],
    ["Device",            "Apple Silicon MPS (torch.device(\"mps\"))"],
]
add_table(rows, col_widths=[5.0, 13.0])

# 학습 곡선 이미지
add_image(
    "/Users/bongjunpyo/smartfarm-upgrade/report_assets/training_curve.png",
    width_cm=16,
    caption="그림 1. EfficientNet-B3 학습 곡선 — 10 epochs / AdamW + CosineLR"
)

add_heading2("4.5 학습 결과 및 평가")
add_para(
    "10 epoch 학습 결과 검증셋 기준 95.30%의 정확도를 달성하였다. Train/Val 격차가 약 0.2%p로 "
    "과적합 없이 일반화 성능이 잘 확보되었음을 확인하였다."
)
rows = [
    ["지표",                "값",       "비고"],
    ["Best Val Accuracy",   "95.30%",   "10 epoch 학습 / 4,000 val 샘플"],
    ["Best Val Loss",       "0.16",     "CrossEntropy (가중 평균)"],
    ["Train Accuracy",      "95.5%",    "동일 epoch 기준 train 정확도"],
    ["Train/Val Gap",       "0.2%p",    "과적합 거의 없음"],
    ["저장 모델",            "best_crop_model.pth",  "PyTorch state_dict"],
    ["ONNX 모델",            "best_crop_model.onnx", "약 41 MB (opset 17)"],
]
add_table(rows, col_widths=[4.5, 4.0, 9.5])

add_heading3("정성 평가 — 실제 추론 결과")
add_image(
    "/Users/bongjunpyo/smartfarm-upgrade/report_assets/inference_gallery.png",
    width_cm=16,
    caption="그림 2. ONNX 추론 결과 8장 (test_images) — 초록 테두리: 건강 / 빨강 테두리: 병해"
)

add_heading2("4.6 ONNX 변환 및 추론 검증")
add_para(
    "PyTorch에서 학습한 모델을 다양한 환경에서 재사용할 수 있도록 ONNX 포맷으로 변환하였다. "
    "변환 후에는 onnxruntime을 사용해 동일 입력에 대한 출력값이 PyTorch와 일치하는지를 "
    "수치 검증하였다."
)
add_para("ONNX 변환은 다음 설정으로 수행하였다.")
add_info_box(
    "torch.onnx.export(\n"
    "    model, dummy_input, \"best_crop_model.onnx\",\n"
    "    input_names=[\"input\"], output_names=[\"output\"],\n"
    "    opset_version=17,\n"
    ")"
)

add_heading2("4.7 그리드 추론 알고리즘")
add_para(
    "단일 추론만으로는 \"잎 전체가 병해\"라는 단일 판정만 가능하다는 한계가 있다. 이를 보완하기 "
    "위해 이미지를 10×10 격자(=100개 셀)로 분할하고 각 셀을 독립적으로 추론하여, 셀 단위 "
    "병해 비율을 산출하는 그리드 추론 알고리즘을 추가로 설계하였다."
)
add_image(
    "/Users/bongjunpyo/smartfarm-upgrade/report_assets/grid_overlay.png",
    width_cm=15,
    caption="그림 3. 10×10 그리드 추론 오버레이 — 빨강: 병해 셀 / 초록: 건강 셀"
)
add_para(
    "그리드 추론은 각 셀의 출력을 (label, confidence) 튜플로 모은 뒤 \"병해 셀 비율\"을 다음 "
    "수식으로 산출한다."
)
add_info_box(
    "Disease Ratio = (Σ I(label_i ≠ Healthy)) / N    where N = 100, I = 지시함수"
)
add_para(
    "예시 이미지(TomatoEarlyBlight1.JPG)에서는 100개 셀 중 49개가 병해로 판정되어 49%의 "
    "위험 비율이 산출되었다. 이 비율은 그대로 모바일 앱의 위험도 표시(NORMAL/WARNING/"
    "DANGER)에 매핑되며, 진단의 \"근거\"를 사용자에게 시각적으로 보여주는 핵심 출력이 된다."
)

# ============================================================
# 5. AI 응용 — 모바일 앱 연동
# ============================================================
add_heading1("5. AI 모델의 응용 — 모바일 앱 연동")
add_para(
    "본 장에서는 학습된 AI 모델이 실제 사용자에게 어떤 형태로 전달되는지를 정리한다. "
    "본 보고서는 인공지능실습 보고서이므로 화면 디자인보다는 \"AI 모델 출력이 어떤 응용 "
    "로직을 거쳐 사용자 화면에 표시되는가\"에 초점을 둔다."
)

add_heading2("5.1 백엔드 추론 API (FastAPI)")
add_para(
    "FastAPI를 이용해 ONNX 모델을 서비스화하였다. /analyze 엔드포인트는 multipart/form-data로 "
    "(field_id, file)을 받아 이미지를 디코드한 뒤 전처리 → onnxruntime 추론 → softmax → "
    "분류 결과를 DB에 기록하고 응답한다."
)
add_image(
    f"{PIC}/KakaoTalk_Photo_2026-06-08-12-09-18.png",
    width_cm=15,
    caption="그림 4. FastAPI Swagger UI — /analyze, /fields, /images 등 9개 엔드포인트"
)
add_image(
    f"{PIC}/KakaoTalk_Photo_2026-06-08-12-09-32.png",
    width_cm=15,
    caption="그림 5. GET /fields 실제 응답 — disease_type=\"RUST\", confidence=0.92로 추론 결과가 영속화됨"
)
add_image(
    f"{PIC}/KakaoTalk_Photo_2026-06-08-12-09-40.png",
    width_cm=15,
    caption="그림 6. POST /analyze 요청 폼 — field_id와 이미지를 보내면 ONNX 추론 후 분류·신뢰도 반환"
)
add_para("주요 API는 다음과 같다.")
rows = [
    ["메서드",  "엔드포인트",                              "기능"],
    ["POST",   "/analyze",                                "이미지 업로드 + AI 추론 + DB 저장 (핵심)"],
    ["GET",    "/fields",                                 "전체 구역 + 최신 분석 결과 조회"],
    ["GET",    "/status/{field_id}",                      "특정 구역의 상태와 신뢰도 조회"],
    ["GET",    "/images",                                 "촬영 이미지 목록 (사진 기록 화면용)"],
    ["GET",    "/history",                                "구역별 분석 이력 시계열"],
    ["GET",    "/notifications",                          "위험 알림 목록"],
    ["PATCH",  "/notifications/{notification_id}/read",   "알림 읽음 처리"],
    ["POST",   "/admin/run-daily-capture",                "하루 1회 촬영 트리거 (테스트용)"],
    ["GET",    "/health",                                 "헬스 체크"],
]
add_table(rows, col_widths=[2.0, 5.5, 10.5])

add_heading2("5.2 안드로이드 앱 화면 구성")
add_para(
    "안드로이드 앱은 백엔드 API로부터 받은 AI 결과를 사용자에게 직관적으로 전달하기 위한 "
    "최종 인터페이스이다. 본 보고서는 AI 출력이 화면에 어떻게 매핑되는지에 초점을 둔다."
)

add_heading3("(1) 스플래시 — 앱 진입")
add_image(
    f"{PIC}/KakaoTalk_Photo_2026-06-08-12-09-08.png",
    width_cm=7.5,
    caption="그림 7. SmartFarm 스플래시 화면 — \"스마트 질병 감지 시스템\""
)

add_heading3("(2) 구역 현황 — AI 분석 결과 요약")
add_para(
    "상단에 \"정상 / 위험 / 주의\" 카운트가 표시되고, 아래 격자에는 각 구역(A1~B4)의 최신 "
    "AI 분석 결과가 카드로 표시된다. AI 모델이 위험으로 판정한 구역은 빨간 테두리와 함께 "
    "신뢰도(예: A4 / 위험 100%)가 함께 표출된다."
)
add_image(
    f"{PIC}/KakaoTalk_Photo_2026-06-08-12-08-33.png",
    width_cm=8.0,
    caption="그림 8. 구역 현황 화면 — 정상 18, 위험 2, 주의 0 / A4 구역 신뢰도 100%로 위험 판정"
)

add_heading3("(3) 구역 상세 — AI 진단 근거")
add_para(
    "구역 카드를 누르면 해당 구역의 최신 촬영 이미지와 AI 분석 결과 상세(상태, 질병명, "
    "신뢰도, 촬영 시각)가 표시된다. AI 출력이 사용자에게 직접 노출되는 화면이며, 신뢰도가 "
    "낮을 때는 사용자가 직접 재촬영을 결정할 수 있도록 수치를 그대로 보여준다."
)
p = d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r1 = p.add_run()
r1.add_picture(f"{PIC}/KakaoTalk_Photo_2026-06-08-12-08-42.png", width=Cm(7.5))
r1 = p.add_run("    ")
r1 = p.add_run()
r1.add_picture(f"{PIC}/KakaoTalk_Photo_2026-06-08-12-08-52.png", width=Cm(7.5))
pc = d.add_paragraph()
pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
rc = pc.add_run("그림 9. (좌) A4 구역 — DANGER / RUST / 신뢰도 100%   (우) B2 구역 — NORMAL / 신뢰도 98%")
set_run_font(rc, size=10, italic=True, color=(0x47, 0x55, 0x69))
pc.paragraph_format.space_after = Pt(10)

add_heading3("(4) 사진 기록 — 일별 누적 진단 이력")
add_para(
    "하루 1회 촬영된 이미지를 날짜별로 누적하여 갤러리 형태로 보여준다. 각 카드 좌상단에 "
    "AI가 판정한 질병 태그(RUST 등)와 신뢰도가 표시되어, 시간이 지남에 따라 동일 구역의 "
    "병해 진행 상태를 한눈에 추적할 수 있다."
)
add_image(
    f"{PIC}/KakaoTalk_Photo_2026-06-08-12-07-58.png",
    width_cm=8.0,
    caption="그림 10. 사진 기록 화면 — RUST 100% / RUST 92% / 정상 카드가 날짜순으로 누적"
)

add_heading3("(5) 알림 — 위험 임계치 초과 시 자동 생성")
add_para(
    "AI 추론 결과 신뢰도가 위험 임계치를 초과하면 백엔드가 자동으로 notifications 테이블에 "
    "알림을 생성하고, 앱은 이를 별도 화면에 노출한다. 즉 AI 출력 → DB 트리거 → 사용자 알림 "
    "흐름이 완전 자동화되어 있다."
)
add_image(
    f"{PIC}/KakaoTalk_Photo_2026-06-08-12-09-00.png",
    width_cm=7.5,
    caption="그림 11. 알림 화면 — [a4] RUST 감지됨 (신뢰도 100%) / 2026-06-04 12:27"
)

# ============================================================
# 6. DB 설계
# ============================================================
add_heading1("6. 데이터베이스 설계")
add_para(
    "AI 추론 결과를 영구적으로 보관하고 이력을 추적할 수 있도록 PostgreSQL을 사용하였다. "
    "fields(구역) – images(이미지) – analysis_results(AI 진단 결과) – notifications(알림)의 "
    "관계형 구조로 설계하여, 어떤 구역의 어떤 이미지에 대해 AI가 어떤 판정을 내렸고, 그 결과로 "
    "어떤 알림이 발생했는지를 모두 역추적할 수 있다."
)
add_image(
    f"{LEG}/image1.png",
    width_cm=15,
    caption="그림 12. PostgreSQL ERD — fields · images · analysis_results · notifications"
)
rows = [
    ["테이블",           "주요 컬럼",                                                                "AI와의 관계"],
    ["fields",           "id, name, location, status, created_at, updated_at",                       "AI 진단 결과(status)가 NORMAL/WARNING/DANGER로 반영됨"],
    ["images",           "id, file_path, file_size_kb, captured_at, created_at",                     "AI 모델의 입력 이미지"],
    ["analysis_results", "id, field_id, image_id, disease_type, confidence, analyzed_at",            "AI 모델의 출력 (클래스 + 신뢰도)을 영속화"],
    ["notifications",    "id, field_id, analysis_id, message, is_read, created_at",                  "AI 신뢰도가 임계치를 넘으면 자동 생성되는 트리거 결과"],
]
add_table(rows, col_widths=[3.5, 6.0, 8.5])

# ============================================================
# 7. 구현 결과 및 검증
# ============================================================
add_heading1("7. 구현 결과 및 검증")
rows = [
    ["영역",                "구현 결과",                                                          "상태"],
    ["AI 모델 학습",        "EfficientNet-B3 / 10 epoch / Val Acc 95.30%",                         "완료"],
    ["AI 모델 변환",        "ONNX (opset 17) 변환 + onnxruntime 추론값 검증",                       "완료"],
    ["전처리 파이프라인",   "HSV Anomaly Score 기반 70×70 크롭 + 224×224 업스케일",                  "완료"],
    ["그리드 추론",         "10×10 셀별 추론 → 병해 비율 산출",                                     "완료"],
    ["백엔드 API",          "/analyze, /fields, /history, /notifications 9개 엔드포인트",            "완료"],
    ["DB 설계",             "fields · images · analysis_results · notifications 4테이블",            "완료"],
    ["안드로이드 앱",       "구역 현황 · 구역 상세 · 사진 기록 · 알림 4화면",                          "완료"],
    ["전체 통합 테스트",    "촬영 → 업로드 → AI 추론 → DB 저장 → 알림 → 앱 표출 전 흐름 확인",        "완료"],
]
add_table(rows, col_widths=[4.0, 11.0, 3.0])

# ============================================================
# 8. 역할 분담
# ============================================================
add_heading1("8. 역할 분담")
add_para("프로젝트는 4인 1조로 진행되었으며, 각 팀원의 담당 영역은 다음과 같다.")

rows = [
    ["팀원",     "직책",   "맡은 역할",                                              "주요 산출물"],
    ["봉준표",   "팀장",   "AI 모델 학습 · Android–AI 연동 · 프로젝트 설계 / 총괄",     "EfficientNet-B3 학습, ONNX 변환, 그리드 추론 로직, PlantDiseaseClassifier.java, 전체 아키텍처 설계"],
    ["이동제",   "팀원",   "백엔드 개발 · 프론트엔드 수정 · DB 수정",                  "FastAPI 추론 라우터, 앱 UI 보정, DB 스키마 수정 및 마이그레이션"],
    ["유광현",   "팀원",   "DB 구성 · FastAPI 연동",                                 "PostgreSQL 4테이블 구성, SQLAlchemy 모델 정의, FastAPI ↔ DB 비동기 연결"],
    ["최성옥",   "팀원",   "앱 개발 · 모니터링 화면 구현",                            "안드로이드 Activity/Fragment, 구역 현황·구역 상세·사진 기록·알림 4화면 구현"],
]
add_table(rows, col_widths=[2.2, 1.8, 5.5, 8.5])

# ============================================================
# 9. 느낀 점
# ============================================================
add_heading1("9. 개발 과정에서 느낀 점")
add_para("인공지능실습 관점에서 본 프로젝트를 진행하며 다음과 같은 점을 배웠다.")

add_heading3("(1) 정확도가 좋아 보이는 모델이 \"좋은 모델\"은 아니다")
add_para(
    "1차 시도로 학습한 YOLOv8s 모델은 mAP50이 0.994까지 나왔지만, 모바일 환경에서 두 모델을 "
    "동시에 로드해야 하는 부담과 추론 지연 때문에 결국 폐기하였다. AI 실습에서 가장 큰 교훈은 "
    "\"수치가 잘 나오는가\"보다 \"실제 응용 환경에 들어갈 수 있는가\"가 더 중요한 평가 기준이라는 "
    "점이다."
)

add_heading3("(2) 모델보다 전처리가 정확도를 좌우한다")
add_para(
    "동일한 EfficientNet-B3을 사용하더라도, 어떤 70×70 타일을 학습 데이터로 선택하느냐에 따라 "
    "정확도가 5%p 이상 변동하였다. HSV Anomaly Score 도입 이후 학습이 안정화되었고, 모델 자체를 "
    "교체하는 것보다 데이터를 선별하는 것이 훨씬 효율적이라는 사실을 체감하였다."
)

add_heading3("(3) PyTorch → ONNX → onnxruntime의 간극")
add_para(
    "torch.onnx.export()는 한 줄로 끝나지만, 변환 후 onnxruntime에서 동일한 출력이 나오는지를 "
    "수치 검증하는 과정에서 입력 텐서 정규화(NHWC↔NCHW), softmax 위치, 데이터 타입 캐스팅 등 "
    "프레임워크 간 경계 작업이 필요했다. \"학습 = 모델 완성\"이 아니라 \"배포 가능한 형태 = 모델 "
    "완성\"이라는 관점을 갖게 되었다."
)

add_heading3("(4) AI 출력의 \"사용자 번역\" 단계가 중요하다")
add_para(
    "EfficientNet-B3의 출력은 23차원 softmax 확률 벡터에 불과하지만, 사용자는 이것을 \"우리 밭에 "
    "병이 있는가 / 얼마나 심각한가\"로 이해해야 한다. 그래서 신뢰도 임계치를 NORMAL / WARNING / "
    "DANGER로 매핑하고, 그리드 추론으로 \"어디에\" 병이 있는지까지 시각화하는 응용 단계가 모델 "
    "자체만큼 중요하다는 것을 알게 되었다."
)

add_heading3("(5) 인터페이스부터 합의한 협업이 통했다")
add_para(
    "AI 팀(봉준표)과 백엔드 팀(이동제·유광현), 앱 팀(최성옥)이 1주차에 Java 클래스의 public 메서드 "
    "시그니처와 FastAPI 응답 JSON 스키마를 먼저 합의했기 때문에, 이후 각자 병렬로 작업하고 "
    "마지막에 한 번에 통합할 수 있었다. AI 실습이 단독 작업이 아닌 풀스택 협업으로 확장될 때 "
    "필요한 가장 큰 자산은 \"잘 정의된 인터페이스\"임을 확인하였다."
)

# ============================================================
# 10. 향후 개선
# ============================================================
add_heading1("10. 향후 개선 계획")
add_bullet("TFLite 변환 및 INT8 양자화를 적용하여 안드로이드 온디바이스 추론 속도를 2배 이상 향상시킨다.")
add_bullet("GradCAM 시각화를 백엔드 API에 정식 응답 필드로 추가하여 \"AI가 어디를 보고 판정했는지\"를 사용자에게 직접 제공한다.")
add_bullet("작물 클래스를 5종(현재) → 10~15종(딸기·상추·오이 등)으로 확장하여 한국 비닐하우스 작물 커버리지를 높인다.")
add_bullet("신뢰도 임계치(현재 고정값)를 농가별 사용 패턴 학습으로 자동 조정하도록 변경한다.")
add_bullet("ONNX 추론 서버 응답 시간을 비동기 큐(Celery + Redis)로 분리해 동시 요청에서도 1초 이내 응답을 유지한다.")
add_bullet("AI 클래스명을 한글/영문 매핑 테이블로 정리하고, 농촌진흥청 데이터와 연계하여 질병별 권장 방제 정보를 자동 추천한다.")

# ============================================================
# 11. 결론
# ============================================================
add_heading1("11. 결론")
add_para(
    "본 프로젝트는 인공지능실습 과목의 목표인 \"실제 데이터로 모델을 학습하고 평가하여 활용 "
    "가능한 시스템으로 연결한다\"를 충실히 수행하였다. Kaggle 공개 데이터셋을 활용해 "
    "EfficientNet-B3 모델을 학습하고 23개 식물 병해 클래스를 95.30%의 정확도로 분류하는 모델을 "
    "확보하였으며, HSV Anomaly Score 기반 전처리와 10×10 그리드 추론이라는 두 가지 응용 알고리즘을 "
    "직접 설계하여 모델의 활용도를 끌어올렸다."
)
add_para(
    "또한 학습된 모델을 ONNX로 변환한 뒤 FastAPI 백엔드 + PostgreSQL DB + 안드로이드 앱과 "
    "통합하여, 단순한 학습 실습 결과물이 아니라 실제 농가 사용자가 사용 가능한 형태의 응용 "
    "서비스로 확장하였다. 본 프로젝트를 통해 데이터 전처리 — 모델 학습 — 평가 — 변환 — 배포 — UI "
    "응용으로 이어지는 AI 파이프라인의 전체 과정을 풀스택으로 경험하였으며, 단일 모델의 정확도가 "
    "아닌 \"전체 시스템의 정확도\"가 인공지능의 실제 가치임을 확인하였다."
)

# 끝페이지 마무리
add_para("", space_after=20)
p = d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
r = p.add_run("— 보고서 끝 —")
set_run_font(r, size=12, italic=True, color=(0x64, 0x74, 0x8b))

p = d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("2026. 06. 08.  /  인공지능실습  /  봉준표 · 이동제 · 유광현 · 최성옥")
set_run_font(r, size=10, color=(0x64, 0x74, 0x8b))

d.save(OUT)
print(f"\nSaved: {OUT}")
import os
print(f"Size: {os.path.getsize(OUT)/1024:.1f} KB")
